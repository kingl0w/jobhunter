import json
import logging
import os
import re
from datetime import datetime

from docx import Document
from google import genai
from google.genai import types
from sqlalchemy.orm import Session

from config import settings
from database import get_resume, update_application_status
from models import ResumeScore

log = logging.getLogger(__name__)

MODEL = "gemini-2.5-flash"
TAILORED_DIR = settings.tailored_resumes_dir
os.makedirs(TAILORED_DIR, exist_ok=True)

_client: genai.Client | None = None


def _get_client() -> genai.Client:
    global _client
    if _client is None:
        _client = genai.Client(api_key=settings.gemini_api_key)
    return _client


def summarize_job(description: str) -> str:
    response = _get_client().models.generate_content(
        model=MODEL,
        contents=(
            "Summarize this job posting in exactly 3 sentences and under 80 words. "
            "Sentence 1: what the role does day-to-day. "
            "Sentence 2: what tech stack and tools are required. "
            "Sentence 3: what kind of company it is. "
            "Return only the summary, no preamble.\n\n"
            f"{description}"
        ),
        config=types.GenerateContentConfig(
            max_output_tokens=400,
            thinking_config=types.ThinkingConfig(thinking_budget=0),
        ),
    )
    return response.text.strip()


def _strip_code_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        parts = text.split("\n", 1)
        text = parts[1] if len(parts) > 1 else text[3:]
    if text.endswith("```"):
        text = text[:-3]
    return text.strip()


def _read_docx(path: str) -> Document:
    return Document(path)


_LABELED_SKILL_RE = re.compile(r"^[A-Z][A-Za-z &/]+:\s")


# TODO: only rewrites the skills section. A future pass could also tailor the
# professional summary or specific experience bullets, but those require much
# more careful prompting to avoid fabricating claims.
def _extract_skills_section(doc: Document) -> tuple[int, int, list[str], str | None]:
    """Return (start, end, skill_lines, pattern) where pattern is A/B/C or None."""
    paragraphs = doc.paragraphs

    # Pattern A: explicit "Skills" heading (or bold line) followed by content.
    a_start = None
    a_end = None
    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        is_heading = para.style.name.lower().startswith("heading")
        is_bold_line = bool(para.runs) and all(run.bold for run in para.runs if run.text.strip())

        if (is_heading or is_bold_line) and "skill" in text:
            a_start = i + 1
            continue

        if a_start is not None and (is_heading or is_bold_line) and text:
            a_end = i
            break

    if a_start is not None:
        if a_end is None:
            a_end = len(paragraphs)
        lines = [paragraphs[i].text for i in range(a_start, a_end) if paragraphs[i].text.strip()]
        if lines:
            log.info("skills section detected: pattern=A, paragraphs %d-%d", a_start, a_end - 1)
            return a_start, a_end, lines, "A"

    # Pattern B: consecutive run of "Category: items, ..." paragraphs with no parent heading.
    best_start = -1
    best_end = -1
    run_start = None
    for i, para in enumerate(paragraphs):
        if _LABELED_SKILL_RE.match(para.text):
            if run_start is None:
                run_start = i
        else:
            if run_start is not None and (i - run_start) >= 2 and (i - run_start) > (best_end - best_start):
                best_start, best_end = run_start, i
            run_start = None
    if run_start is not None and (len(paragraphs) - run_start) >= 2 and (len(paragraphs) - run_start) > (best_end - best_start):
        best_start, best_end = run_start, len(paragraphs)

    if best_start != -1:
        lines = [paragraphs[i].text for i in range(best_start, best_end) if paragraphs[i].text.strip()]
        log.info("skills section detected: pattern=B, paragraphs %d-%d", best_start, best_end - 1)
        return best_start, best_end, lines, "B"

    # Pattern C: a single flat comma-list paragraph near the top, before the first heading.
    first_heading = None
    for i, para in enumerate(paragraphs):
        if para.style.name.lower().startswith("heading") and para.text.strip():
            first_heading = i
            break
    scan_end = first_heading if first_heading is not None else min(len(paragraphs), 15)
    for i in range(2, scan_end):
        text = paragraphs[i].text.strip()
        if not text:
            continue
        tokens = [t.strip() for t in text.split(",") if t.strip()]
        if len(tokens) >= 3 and sum(len(t) for t in tokens) / len(tokens) < 30:
            log.info("skills section detected: pattern=C, paragraph %d", i)
            return i, i + 1, [text], "C"

    return -1, -1, [], None


def _replace_skills_section(doc: Document, start: int, end: int, new_lines: list[str]) -> Document:
    paragraphs = doc.paragraphs

    first_skill_para = paragraphs[start] if start < len(paragraphs) else None
    template_run = None
    if first_skill_para and first_skill_para.runs:
        template_run = first_skill_para.runs[0]

    for i in range(end - 1, start - 1, -1):
        para_element = paragraphs[i]._element
        para_element.getparent().remove(para_element)

    insert_point = paragraphs[start - 1]._element if start > 0 else doc.element.body[0]

    for line in reversed(new_lines):
        new_para = doc.add_paragraph()
        new_run = new_para.add_run(line)

        if template_run:
            new_run.bold = template_run.bold
            new_run.italic = template_run.italic
            new_run.underline = template_run.underline
            if template_run.font.size:
                new_run.font.size = template_run.font.size
            if template_run.font.name:
                new_run.font.name = template_run.font.name

        new_element = new_para._element
        new_element.getparent().remove(new_element)
        insert_point.addnext(new_element)

    return doc


def tailor_resume(
    job_id: str,
    description: str,
    resume_id: str,
    db: Session,
) -> str:
    resume = get_resume(db, resume_id)
    if not resume:
        raise ValueError(f"resume not found: {resume_id}")

    base_path = resume.file_path
    if not os.path.isfile(base_path):
        log.error("base resume file missing: %s", base_path)
        raise FileNotFoundError(base_path)

    if not base_path.lower().endswith(".docx"):
        log.warning("tailoring only supported for .docx, returning base: %s", base_path)
        return base_path

    score_row = (
        db.query(ResumeScore)
        .filter(ResumeScore.job_id == job_id, ResumeScore.resume_id == resume_id)
        .first()
    )
    missing_keywords = (score_row.missing_keywords if score_row else []) or []

    doc = _read_docx(base_path)
    start, end, current_skill_lines, _pattern = _extract_skills_section(doc)

    if start == -1 or not current_skill_lines:
        log.warning("no skills section found in %s, returning base resume", base_path)
        return base_path

    prompt = (
        "You are a resume tailoring assistant. Your job is to subtly update the "
        "skills section of a resume so it better matches a job posting.\n\n"
        "RULES:\n"
        "- Only modify the skill keyword lines below.\n"
        "- Incorporate the missing keywords where they fit naturally.\n"
        "- Do NOT invent skills the candidate doesn't have — only reword, "
        "reorder, or add keywords that are reasonable synonyms or related.\n"
        "- Do NOT change job titles, company names, dates, or metrics.\n"
        "- Keep the same number of lines and similar line lengths.\n"
        "- Return ONLY a JSON object with key \"updated_skills_lines\" containing "
        "a list of strings, one per skill category line.\n\n"
        f"JOB DESCRIPTION:\n{description}\n\n"
        f"CURRENT SKILLS LINES:\n{json.dumps(current_skill_lines)}\n\n"
        f"MISSING KEYWORDS TO INCORPORATE:\n{json.dumps(missing_keywords)}\n\n"
        "Respond with only the JSON object, no markdown fences or extra text."
    )

    response = _get_client().models.generate_content(
        model=MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            max_output_tokens=4096,
            thinking_config=types.ThinkingConfig(thinking_budget=0),
        ),
    )
    raw_text = _strip_code_fences(response.text.strip())

    try:
        parsed = json.loads(raw_text)
        updated_lines = parsed["updated_skills_lines"]
    except (json.JSONDecodeError, KeyError, TypeError) as exc:
        log.error("failed to parse gemini response: %s\nraw: %s", exc, raw_text[:500])
        return base_path

    if not isinstance(updated_lines, list) or not all(isinstance(s, str) for s in updated_lines):
        log.error("unexpected response shape: %s", type(updated_lines))
        return base_path

    tailored_doc = _read_docx(base_path)
    _replace_skills_section(tailored_doc, start, end, updated_lines)

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    short_id = job_id[:12]
    label_slug = "".join(c if c.isalnum() else "_" for c in (resume.label or "resume")).strip("_")[:40]
    filename = f"{short_id}_{label_slug}_{timestamp}.docx"
    output_path = os.path.join(TAILORED_DIR, filename)
    tailored_doc.save(output_path)

    log.info("tailored resume saved: %s", output_path)

    update_application_status(db, job_id, resume_used=output_path)

    return output_path
